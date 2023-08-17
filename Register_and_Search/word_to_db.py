import sqlite3
import re
import docx
import os
#import datetime
#import pytz


Pattern_T1 = re.compile(r"^(■「)(.*)(」)$")
Pattern_T2 = re.compile(r"^(【)(.*)(】)$")
Pattern_P1 = re.compile(r"(■場所・)(.*)")
Pattern_D1 = re.compile(r"^(.{2})(　)(.*)")
Pattern_D2 = re.compile(r"^(　{3})(.*)")
Pattern_D3 = re.compile(r"^(.{3})(　)(.*)")

class Word_to_DB:
    #指定したワードファイルを読み込む(メタ情報も一緒に入力)
    def __init__(self, filepath, filename, created_date, character_list):
        self.doc = docx.Document(filepath)
        self.filepath = os.path.basename(filepath)
        self.state = 0
        self.filename = filename
        self.flag = False
        self.count = 1
        self.date = created_date
        self.data_list = []
        self.C_Master_Table = character_list
        self.character_dict = {}
        self.input_type = 0

    def get_chara_name(self, chara):
        chara = chara.replace("　", "")
        chara_name = ""
        for i in self.C_Master_Table:
            if chara in i:
                chara_name = i
        if not len(chara_name) == 0:
            self.character_dict[chara] = chara_name
        else:
            self.character_dict[chara] = chara
            chara_name = chara
        return chara_name
    
    def append_data(self, chara, dialogue):
        dialogue = dialogue.replace("「", "")
        dialogue = dialogue.replace("」", "")
        current_data = [self.filename, chara, dialogue, self.count, self.place, self.date, self.filepath]
        self.data_list.append(current_data)
        self.count +=  1
    
    def write_db(self):
        print("データ書き込み中")
        conn = sqlite3.connect("db.sqlite3")
        cursor = conn.cursor()
        sql = "INSERT INTO Register_and_Search_dialogue (filename, character, dialogue, dialogue_number, place, created_date, filepath) VALUES (?, ?, ?, ?, ?, ?, ?)"
        cursor.executemany(sql, self.data_list)
        conn.commit()
        conn.close()
        print("データ書き込み完了")

    def input_type_1(self):
        text = self.doc.paragraphs
        for p in text:
            current_line = p.text
            if result := Pattern_P1.search(current_line):
                self.place = result.group(2)
                self.flag = True

            elif self.flag == True:
                # stateが0の場合
                if self.state == 0:
                    if result := Pattern_D2.search(current_line):
                        chara_name = "ト書き"
                        dialogue = result.group(2)
                        self.state = 2
                    elif result := Pattern_D1.search(current_line):
                        chara = result.group(1).replace("　", "")
                        if chara in self.character_dict:
                            chara_name = self.character_dict[chara]
                        else:
                            chara_name = self.get_chara_name(chara)
                        dialogue = result.group(3)
                        self.state = 1
                    else:
                        pass

                elif self.state == 1:
                    if len(current_line) == 0:
                        self.append_data(chara_name, dialogue)
                        self.state = 0

                    elif result := Pattern_D2.search(current_line):
                        dialogue = dialogue + result.group(2)

                    elif result := Pattern_D1.search(current_line):
                        self.append_data(chara_name, dialogue)
                        chara = result.group(1).replace("　", "")
                        if chara in self.character_dict:
                            chara_name = self.character_dict[chara]
                        else:
                            chara_name = self.get_chara_name(chara)
                        dialogue = result.group(3)                    
                    else:
                        dialogue = dialogue + current_line
                
                elif self.state == 2:
                    if len(current_line) == 0:
                        self.append_data(chara_name, dialogue)
                        self.state = 0
                    else:
                        dialogue = dialogue + current_line

        if not self.state == 0:
            self.append_data(chara_name, dialogue)
        self.write_db()

    def input_type_2(self):
        text = self.doc.paragraphs
        self.place = None
        for p in text:
            current_line = p.text
            if self.state == 0:
                if result := Pattern_D3.search(current_line):
                    chara = result.group(1)
                    if chara in self.character_dict:
                        chara_name = self.character_dict[chara]
                    else:
                        chara_name = self.get_chara_name(chara)
                    dialogue = result.group(3)
                    self.state = 1
                else:
                    pass
            
            elif self.state == 1:
                current_line = current_line.replace("　", "")
                if len(current_line) == 0:
                    self.append_data(chara_name, dialogue)
                    self.state = 0
                else:
                    dialogue = dialogue + current_line
        if self.state == 1:
            self.append_data(chara_name, dialogue)
        self.write_db()

    def clean_space(self, input_text):
        output_text = input_text.replace(" ", "")
        output_text = output_text.replace("　", "")
        return output_text

    def run(self):
        text = self.doc.paragraphs
        first_line = text[0].text
        second_line = text[1].text
        first_line = self.clean_space(first_line)
        if result := Pattern_T1.search(first_line):
            second_line = self.clean_space(second_line)
            if len(second_line) == 0:
                self.input_type_1()
                self.input_type = 1
            else:
                print("非対応のフォーマットです")
        elif result := Pattern_T2.search(first_line):
            second_line = self.clean_space(second_line)
            if len(second_line) == 0:
                self.input_type_2()
                self.input_type = 2
            else:
                print("非対応のフォーマットです")
        else:
            print("非対応のフォーマットです")
        return self.input_type
     
"""
if __name__ == "__main__":
    #処理を実行
    chara_list = ["帝 ナギ", "鳳暎一", "鳳瑛二", "美風 藍","聖川 真斗", "来栖 翔"]
    utc_now = datetime.datetime.utcnow()
    jst = pytz.timezone('Asia/Tokyo')
    jst_now = utc_now.astimezone(jst)
    word_path = str(input())
    filename = os.path.basename(word_path)
    run = Word_to_DB(word_path, filename, jst_now, chara_list)
    flag = run.run()
    print(flag)
"""