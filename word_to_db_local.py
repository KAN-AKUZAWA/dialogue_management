import sqlite3
import re
import docx
import os
import datetime
import pytz
import glob
import hashlib
import shutil

# フォーマットごとにテキスト判定パターンを定義
Pattern_T1 = re.compile(r"^(■「)(.*)(」)$")
Pattern_T2 = re.compile(r"^(【)(.*)(】)$")
Pattern_P1 = re.compile(r"(■場所・)(.*)")
Pattern_D1 = re.compile(r"^(.{2})(　)(.*)")
Pattern_D2 = re.compile(r"^(　{3})(.*)")
Pattern_D3 = re.compile(r"^(.{3})(　)(.*)")

class Word_to_DB:
    #指定したワードファイルを読み込む(メタ情報も一緒に入力)
    def __init__(self, filepath, filename, created_date, character_list):
        self.doc = docx.Document(filepath)
        self.filepath = os.path.basename(filepath)
        self.state = 0
        self.filename = filename
        self.flag = False
        self.count = 1
        self.date = created_date
        self.data_list = []
        self.C_Master_Table = character_list
        self.character_dict = {}
        self.input_type = 0
        self.local_path = filepath
        self.subfolder = os.path.basename(os.path.dirname(filepath))

    def move_file(self, source_path, destination_path):
        try:
            shutil.move(source_path, destination_path)
            print(f"ファイルを移動しました: {source_path} -> {destination_path}")
        except FileNotFoundError:
            print("移動元のファイルが見つかりませんでした。")
        except shutil.Error:
            print("ファイルの移動に失敗しました。")

    def cal_hash(self):
        with open(self.local_path, 'rb') as file:
            hash_object = hashlib.sha256(file.read())
            file_hash = hash_object.hexdigest()
            return file_hash
    
    # 日時情報を日本時間に調整するための関数
    def get_japan_current_time(self):
        # 日本のタイムゾーンを指定
        japan_tz = pytz.timezone('Asia/Tokyo')
        utc_now = datetime.datetime.utcnow()
        japan_now = utc_now.replace(tzinfo=pytz.utc).astimezone(japan_tz)
        return japan_now
    
    # Wordファイルの内容を読み込む関数
    def get_file_content(self):
        with open(self.local_path, 'rb') as file:
            return file.read()

    # DBのテーブル"Register_and_Search_FileUpload"に登録するデータを整形する関数
    def organize_file_data(self):
        uploaded_at = self.get_japan_current_time()
        file_hash = self.cal_hash()
        created_at = self.date
        file_name = self.filename
        file_content = self.get_file_content()
        return [[uploaded_at, file_hash, created_at, file_name, file_content]]

    # DBのテーブル"Register_and_Search_FileUpload"にデータを書き込む関数
    def register_file(self):
        print("ファイルの登録開始")
        data_list = self.organize_file_data()
        conn = sqlite3.connect("db.sqlite3")
        cursor = conn.cursor()
        sql = "INSERT INTO Register_and_Search_FileUpload (uploaded_at, file_hash, created_at, file_name, uploaded_file) VALUES (?, ?, ?, ?, ?)"
        cursor.executemany(sql, data_list)
        conn.commit()
        conn.close()
        print("ファイルの登録完了")

    # シナリオに記載されたキャラ名を基にマスターテーブルを参照して、キャラの正式名称を返す関数
    # プロセス中に一度参照したキャラの正式名称は、辞書に格納することで、マスターテーブルを参照する回数を減らして処理を軽量化
    def get_chara_name(self, chara):
        chara = self.clean_space(chara)
        chara_name = ""
        for i in self.C_Master_Table:
            if chara in i:
                chara_name = i
        if not len(chara_name) == 0:
            self.character_dict[chara] = chara_name
        else:
            self.character_dict[chara] = chara
            chara_name = chara
        return chara_name
    
    # 1つのセリフを1レコードにまとめて、リストに格納する関数
    def append_data(self, chara, dialogue):
        dialogue = dialogue.replace("「", "")
        dialogue = dialogue.replace("」", "")
        current_data = [self.filename, chara, dialogue, self.count, self.place, self.date, self.filepath]
        # 変数"self.data_list"にレコードを格納する(2次元リスト形式)
        self.data_list.append(current_data)
        # セリフ番号のカウントアップ
        self.count +=  1
    
    # 変数"self.data_list"に格納されたデータをDBに書き込む関数
    def write_db(self):
        print("データ書き込み中")
        conn = sqlite3.connect("db.sqlite3")
        cursor = conn.cursor()
        sql = "INSERT INTO Register_and_Search_dialogue (filename, character, dialogue, dialogue_number, place, created_date, filepath) VALUES (?, ?, ?, ?, ?, ?, ?)"
        cursor.executemany(sql, self.data_list)
        conn.commit()
        conn.close()
        print("データ書き込み完了")

    # フォーマット1に対する処理
    def input_type_1(self):
        # ファイルを読み込み、テキストを１行ずつ取り出して処理
        text = self.doc.paragraphs
        self.place = None
        for p in text:
            current_line = p.text
            # テキストに「■場所」が含まれる場合、変数"self.place"を更新する
            if result := Pattern_P1.search(current_line):
                self.place = result.group(2)
                # 「■場所」以降にシナリオが記載されるため、最初の「■場所」を検知したタイミングでFlagをTrueにする
                self.flag = True
            elif self.flag == True:
                # 判定パターン、及び判定後の処理は仕様書参照
                if self.state == 0:
                    if result := Pattern_D2.search(current_line):
                        chara_name = "ト書き"
                        dialogue = result.group(2)
                        self.state = 2
                    elif result := Pattern_D1.search(current_line):
                        chara = result.group(1).replace("　", "")
                        if chara in self.character_dict:
                            chara_name = self.character_dict[chara]
                        else:
                            chara_name = self.get_chara_name(chara)
                        dialogue = result.group(3)
                        self.state = 1
                    else:
                        pass

                elif self.state == 1:
                    if len(current_line) == 0:
                        self.append_data(chara_name, dialogue)
                        self.state = 0
                    elif result := Pattern_D2.search(current_line):
                        dialogue = dialogue + result.group(2)

                    elif result := Pattern_D1.search(current_line):
                        self.append_data(chara_name, dialogue)
                        chara = result.group(1).replace("　", "")
                        if chara in self.character_dict:
                            chara_name = self.character_dict[chara]
                        else:
                            chara_name = self.get_chara_name(chara)
                        dialogue = result.group(3)                    
                    else:
                        dialogue = dialogue + current_line
                
                elif self.state == 2:
                    if len(current_line) == 0:
                        self.append_data(chara_name, dialogue)
                        self.state = 0
                    else:
                        dialogue = dialogue + current_line
        if not self.state == 0:
            self.append_data(chara_name, dialogue)
        self.write_db()

    # フォーマット2に対する処理
    def input_type_2(self):
        # ファイルを読み込み、テキストを１行ずつ取り出して処理
        text = self.doc.paragraphs
        self.place = None
        for p in text:
            current_line = p.text
            # 判定パターン、及び判定後の処理は仕様書参照
            if self.state == 0:
                if result := Pattern_D3.search(current_line):
                    chara = result.group(1)
                    if chara in self.character_dict:
                        chara_name = self.character_dict[chara]
                    else:
                        chara_name = self.get_chara_name(chara)
                    dialogue = result.group(3)
                    self.state = 1
                else:
                    pass
            
            elif self.state == 1:
                current_line = current_line.replace("　", "")
                if len(current_line) == 0:
                    self.append_data(chara_name, dialogue)
                    self.state = 0
                else:
                    dialogue = dialogue + current_line
        if self.state == 1:
            self.append_data(chara_name, dialogue)
        self.write_db()

    # 文字列に含まれるスペースを除去する関数
    def clean_space(self, input_text):
        output_text = input_text.replace(" ", "")
        output_text = output_text.replace("　", "")
        return output_text

    # 想定外のフォーマットとして判定されたファイルに対する処理を定義した関数
    # unsupported_formatフォルダ下に、作成日をフォルダ名としたサブフォルダを作成し、ファイルを移動させる
    def unssuported(self):
        print(self.local_path)
        print("非対応のフォーマットです")
        if not os.path.exists(f"unsupported_format/{self.subfolder}"):
            os.mkdir(f"unsupported_format/{self.subfolder}")
        self.move_file(self.local_path, f"unsupported_format/{self.subfolder}/{self.filename}")

    # インプットファイルのフォーマットを判定するための関数
    def judge_format(self, text):
        flag = "0"
        # Wordのテキストを1行ずつ取り出す
        for index, p in enumerate(text):
            current_line = p.text
            # テキスト中のスペースを除去
            current_line = self.clean_space(current_line)
            # 1行目にのみ実行する処理
            # 事前に定義したパターンにマッチするかを判定する
            if index == 0:
                # "■「"から始まり。"」"で終わる場合は、フォーマット1候補として、後続の処理に繋げる
                if result := Pattern_T1.search(current_line):
                    flag = "1-1"
                # "【"から始まり。"】"で終わる場合は、フォーマット2候補として、後続の処理に繋げる
                elif result := Pattern_T2.search(current_line):
                    flag = "2-1"
                # 上記のパターンにマッチしない場合、想定外のフォーマットであるものと判定
                else:
                    return 0
            # フォーマット1候補に対する処理
            # 文書中に「◆概要」、と「◆出演」を含む場合はフォーマット1と判定する
            if flag == "1-1":
                if current_line == "◆概要":
                    flag = "1-2"
            elif flag == "1-2":
                if current_line == "◆出演":
                    return 1
            # フォーマット2候補に対する処理
            # 文書中に「／／END」を含む場合はフォーマット2と判定する
            elif flag == "2-1":
                 if current_line == "／／END":
                        return 2
        # 1行目はパターンにマッチしたものの、それ以降の判定ではじかれた場合は、想定外のフォーマットとして判定
        return 0
                          
    # プログラムの処理を開始
    def run(self):
        text = self.doc.paragraphs
        self.input_type = self.judge_format(text)

        if self.input_type == 1:
            # ファイルを、テーブル"Register_and_Search_fileupload"に登録する
            self.register_file()
            # テキストを解析し、テーブル"Register_and_Search_dialogue"にデータを追加する
            self.input_type_1()
            # ファイルをuploadsフォルダへ移動する
            self.move_file(self.local_path, f"uploads/{self.filename}")
    
        elif self.input_type == 2:
            # ファイルを、テーブル"Register_and_Search_fileupload"に登録する
            self.register_file()
            # テキストを解析し、テーブル"Register_and_Search_dialogue"にデータを追加する
            self.input_type_2()
            # ファイルをuploadsフォルダへ移動する
            self.move_file(self.local_path, f"uploads/{self.filename}")

        #想定外のフォーマットに対する処理
        elif self.input_type == 0:
            self.unssuported()
        return self.input_type
     
# inputフォルダ内の全ての.docxファイルに対して処理を実行
if __name__ == "__main__":
    # データベースから指定したカラムに含まれるデータを取得する関数
    def get_column_values(db_file, table_name, column_name):
        conn = sqlite3.connect(db_file)
        cursor = conn.cursor()
        cursor.execute(f"SELECT {column_name} FROM {table_name}")
        column_values = [row[0] for row in cursor.fetchall()]
        conn.close()
        return column_values
    # 文字列形式で記載された日付を、日付データに変換する関数
    def str_to_datetime(date_str, date_format="%Y-%m-%d"):
        jst = pytz.timezone('Asia/Tokyo')
        naive_date_obj = datetime.datetime.strptime(date_str, date_format)
        jst_date_obj = jst.localize(naive_date_obj)
        return jst_date_obj
    
    db_file = "db.sqlite3"
    table_name = "Register_and_Search_character"
    column_name = "character"
    input_path = "input/**/*.docx"
    # inputフォルダ無いの全ての.docxファイルのパスを取得
    path_list = glob.glob(input_path)
    # キャラクター名の正式名称のリストを、データベースから取得
    chara_list = get_column_values(db_file, table_name, column_name)
    # .docxファイルを1つずつ取り出して、処理を実行
    for path in path_list:
        filename = os.path.basename(path)
        folder_path = os.path.dirname(path)
        folder_name = os.path.basename(folder_path)
        created_date = str_to_datetime(folder_name)
        run = Word_to_DB(path, filename, created_date, chara_list)
        flag = run.run()